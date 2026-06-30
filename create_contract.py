import docx
doc = docx.Document()
doc.add_paragraph("This Contract is made between Company A and Company B on 2024-01-01.")
doc.add_paragraph("The term of this contract is 2 years. The governing law is California.")
doc.add_paragraph("Payment terms: Net 30. Confidentiality is required.")
doc.add_paragraph("Indemnity: Company A shall indemnify Company B against any claims.")
doc.add_paragraph("Liability Cap: The liability is capped at 100000 USD.")
doc.add_paragraph("Signatories: John Doe, Jane Smith.")
doc.save("contract.docx")
